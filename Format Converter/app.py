from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import os, io, tempfile, markdown, pdfkit, pypandoc
from docx import Document
from bs4 import BeautifulSoup
from pdfminer.high_level import extract_text
from markdownify import markdownify as mdify

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For development only. Adjust in production.
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SUPPORTED_FORMATS = ["markdown", "html", "pdf", "docx", "json"]

def detect_format(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    ext_map = {
        '.md': 'markdown', '.markdown': 'markdown',
        '.html': 'html', '.htm': 'html',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.json': 'json'
    }
    return ext_map.get(ext, '')

def safe_convert(func):
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"[{func.__name__}] Conversion failed: {str(e)}")
    return wrapper

@safe_convert
def md_to_html(md_text: str) -> bytes:
    return markdown.markdown(md_text).encode()

@safe_convert
def md_to_docx(md_text: str) -> bytes:
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        output_path = tmp_file.name
    try:
        pypandoc.convert_text(md_text, 'docx', format='md', outputfile=output_path)
        with open(output_path, 'rb') as f:
            docx_bytes = f.read()
    finally:
        os.remove(output_path)
    return docx_bytes

@safe_convert
def md_to_json(md_text: str) -> dict:
    segments = [seg.strip() for seg in md_text.split('\n\n') if seg.strip()]
    return {"segments": segments}

@safe_convert
def md_to_pdf(md_text: str) -> bytes:
    html = markdown.markdown(md_text)
    config = pdfkit.configuration(wkhtmltopdf="C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe")
    return pdfkit.from_string(html, False, configuration=config)

@safe_convert
def html_to_md(html: str) -> bytes:
    return mdify(html).encode()

@safe_convert
def html_to_pdf(html: str) -> bytes:
    html = html[html.find('<body>'):]
    new_header = '''<!DOCTYPE html>
    <html lang="en">
    <head>
    <meta charset="utf-8"/>
    </head>'''
    html = new_header + html
    return pdfkit.from_string(html, False)

@safe_convert
def html_to_docx(html: str) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    def add_heading(text, level):
        doc.add_heading(text, level=level)

    def add_paragraph(text):
        doc.add_paragraph(text)

    def add_list(items, ordered=False):
        for item in items:
            doc.add_paragraph(item, style='List Number' if ordered else 'List Bullet')

    def add_table(table_tag):
        rows = table_tag.find_all('tr')
        if not rows:
            return
        cols = rows[0].find_all(['td', 'th'])
        table = doc.add_table(rows=len(rows), cols=len(cols))
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                table.cell(i, j).text = cell.get_text(strip=True)

    body = soup.body or soup
    for elem in body.find_all(recursive=False):
        try:
            if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                add_heading(elem.get_text(), level=int(elem.name[1]))
            elif elem.name == 'p':
                add_paragraph(elem.get_text())
            elif elem.name in ['ul', 'ol']:
                items = [li.get_text() for li in elem.find_all('li')]
                add_list(items, ordered=elem.name == 'ol')
            elif elem.name == 'table':
                add_table(elem)
            elif elem.name == 'blockquote':
                doc.add_paragraph(elem.get_text(), style="Intense Quote")
            elif elem.name:
                add_paragraph(elem.get_text())
        except Exception:
            continue

    if not soup.body:
        add_paragraph(soup.get_text())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

@safe_convert
def docx_to_md(docx_bytes: bytes) -> bytes:
    temp_path = "temp.docx"
    with open(temp_path, "wb") as f:
        f.write(docx_bytes)
    md = pypandoc.convert_file(temp_path, 'md', format='docx')
    os.remove(temp_path)
    return md.encode()

@safe_convert
def pdf_to_json(pdf_bytes: bytes) -> dict:
    temp_path = "temp.pdf"
    with open(temp_path, "wb") as f:
        f.write(pdf_bytes)
    text = extract_text(temp_path)
    os.remove(temp_path)
    return {"text": text}

@safe_convert
def pdf_to_md(pdf_bytes: bytes) -> bytes:
    temp_path = "temp.pdf"
    with open(temp_path, "wb") as f:
        f.write(pdf_bytes)
    text = extract_text(temp_path)
    os.remove(temp_path)
    md = "\n\n".join(text.split('\n'))
    return md.encode()

CONVERSION_MAP = {
    ('markdown', 'html'): md_to_html,
    ('markdown', 'pdf'): md_to_pdf,
    ('markdown', 'docx'): md_to_docx,
    ('markdown', 'json'): md_to_json,
    ('html', 'markdown'): html_to_md,
    ('html', 'pdf'): html_to_pdf,
    ('html', 'docx'): html_to_docx,
    ('docx', 'markdown'): docx_to_md,
    ('pdf', 'json'): pdf_to_json,
    ('pdf', 'markdown'): pdf_to_md,
}

@app.post("/convert")
async def convert_document(
    file: UploadFile = File(...),
    target_format: str = Form(...),
    download: bool = Form(False),
):
    try:
        if target_format not in SUPPORTED_FORMATS:
            raise HTTPException(400, "Unsupported target format")

        source_format = detect_format(file.filename)
        if not source_format:
            raise HTTPException(400, "Cannot detect source format")

        converter = CONVERSION_MAP.get((source_format, target_format))
        if not converter:
            raise HTTPException(501, f"Conversion from {source_format} to {target_format} not supported")

        input_bytes = await file.read()
        if source_format in ['markdown', 'html']:
            result = converter(input_bytes.decode())
        else:
            result = converter(input_bytes)
        
        if target_format == 'json':
            return JSONResponse(result)
        
        mime_types = {
            'markdown': 'text/markdown',
            'html': 'text/html',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        if download:
            suffix_map = {
                'markdown': '.md',
                'html': '.html',
                'pdf': '.pdf',
                'docx': '.docx',
            }
            suffix = suffix_map.get(target_format, '')
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(result)
                    tmp_path = tmp.name
            except Exception as e:
                raise HTTPException(500, f"Temporary file creation error: {e}")
            return FileResponse(
                tmp_path,
                media_type=mime_types.get(target_format, 'application/octet-stream'),
                filename=f"converted{suffix}",
                background=None
            )
        media_type = mime_types.get(target_format, 'application/octet-stream')
        return StreamingResponse(io.BytesIO(result), media_type=media_type)
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(500, f"Conversion error: {e}")
